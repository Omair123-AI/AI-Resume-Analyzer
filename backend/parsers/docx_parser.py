"""
docx_parser.py - Extract raw text from DOCX resumes
"""
from docx import Document


def extract_text_from_docx(filepath: str) -> str:
    """Extract all paragraphs and table cells from a DOCX file."""
    try:
        doc = Document(filepath)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables (some resumes use table-based layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX Parser] Error: {e}")
        return ""
